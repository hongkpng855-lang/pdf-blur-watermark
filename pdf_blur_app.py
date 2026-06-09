#!/usr/bin/env python3
"""
PDF Blur + Watermark Tool — Web App
=====================================
Upload a PDF, blur regions, add watermarks, download the result.
"""
import os, uuid, io, base64, math, json as _json
from flask import Flask, request, jsonify, send_file
from PIL import Image, ImageFilter, ImageDraw, ImageFont
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = os.urandom(24)

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
FONT_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
if not os.path.exists(FONT_PATH):
    FONT_PATH = None


def get_session_paths(sid):
    """Get file paths for a session, supporting multi-worker gunicorn."""
    sess_dir = os.path.join(UPLOAD_FOLDER, sid)
    os.makedirs(sess_dir, exist_ok=True)
    return {
        'pdf': os.path.join(sess_dir, 'input.pdf'),
        'originals': os.path.join(sess_dir, 'originals'),
        'processed_dir': os.path.join(sess_dir, 'processed'),
        'meta': os.path.join(sess_dir, 'meta.json'),
    }


def pdf_to_images(pdf_path):
    doc = fitz.open(pdf_path)
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    doc.close()
    return images


def images_to_pdf(images, output_path):
    if not images: return
    if len(images) == 1:
        images[0].save(output_path, "PDF", resolution=150)
    else:
        images[0].save(output_path, "PDF", resolution=150, save_all=True,
                       append_images=images[1:])


def apply_blur(page_img, blur_regions, radius=25):
    img = page_img.copy()
    for region in blur_regions:
        x, y, w, h = region['x'], region['y'], region['w'], region['h']
        if w < 5 or h < 5: continue
        crop = img.crop((x, y, x + w, y + h))
        blurred = crop.filter(ImageFilter.GaussianBlur(radius=radius))
        img.paste(blurred, (x, y))
    return img


def apply_watermark(page_img, wm_config):
    """Apply watermark text tiled diagonally at max size."""
    if not wm_config or not wm_config.get('text'):
        return page_img

    img = page_img.copy()
    txt = wm_config['text']
    opacity = wm_config.get('opacity', 30) / 100.0
    user_size = wm_config.get('size', 80)
    spacing = wm_config.get('spacing', 2.0)
    color = wm_config.get('color', '#888888')

    # Parse color
    color = color.lstrip('#')
    r, g, b = int(color[0:2], 16) if len(color) >= 2 else 0, \
              int(color[2:4], 16) if len(color) >= 4 else 0, \
              int(color[4:6], 16) if len(color) >= 6 else 0

    overlay = Image.new('RGBA', img.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)

    font_size = min(user_size, img.width // 2)
    try:
        font = ImageFont.truetype(FONT_PATH, font_size) if FONT_PATH else ImageFont.load_default()
    except:
        font = ImageFont.load_default()

    bbox = draw.textbbox((0, 0), txt, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    if tw < 10:
        font_size = img.width // 2
        try:
            font = ImageFont.truetype(FONT_PATH, font_size) if FONT_PATH else ImageFont.load_default()
        except:
            font = ImageFont.load_default()
        bbox = draw.textbbox((0, 0), txt, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]

    # Tiled diagonally (-30 deg) across the whole page
    angle = -30
    step_x = int(tw * spacing * 0.9)
    step_y = int(th * spacing * 1.25)

    for ty in range(-step_y, img.height + step_y, step_y):
        for tx in range(-step_x, img.width + step_x, step_x):
            txt_overlay = Image.new('RGBA', (step_x * 2, step_y), (0, 0, 0, 0))
            td = ImageDraw.Draw(txt_overlay)
            td.text((0, 0), txt, font=font, fill=(r, g, b, int(255 * opacity)))
            rotated = txt_overlay.rotate(angle, expand=True, center=(step_x, step_y // 2))
            overlay.paste(rotated, (tx, ty), rotated)

    return Image.alpha_composite(img.convert('RGBA'), overlay).convert('RGB')


# ══════════════════════════════════════════════════════════════
# HTML (embedded)
# ══════════════════════════════════════════════════════════════
HTML = r'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>PDF Tools</title>
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body {
  font-family: 'Segoe UI', system-ui, sans-serif;
  background: #0d0d12; color: #e0e0f0;
  height: 100vh; display: flex; flex-direction: column;
}
.header {
  background: #15151e;
  border-bottom: 1px solid #2a2a3a;
}
.tab-bar {
  display: flex; border-bottom: 1px solid #2a2a3a;
}
.tab-bar .tab {
  padding: 8px 18px; font-size: 13px; cursor: pointer;
  color: #8888a0; border-bottom: 2px solid transparent;
  transition: all 0.15s; user-select: none;
}
.tab-bar .tab:hover { color: #e0e0f0; }
.tab-bar .tab.active {
  color: #6c5ce7; border-bottom-color: #6c5ce7;
}
.toolbar {
  padding: 8px 16px; display: flex; align-items: center;
  justify-content: space-between; gap: 10px; flex-wrap: wrap;
}
.toolbar-left { display: flex; align-items: center; gap: 6px; }
.toolbar-right { display: flex; align-items: center; gap: 6px; }

.btn, .toolbar-right button, .toolbar-right label {
  padding: 6px 12px; border-radius: 5px; border: 1px solid #2a2a3a;
  background: #1c1c2a; color: #e0e0f0; cursor: pointer; font-size: 12px;
}
.toolbar-right button:hover, .toolbar-right label:hover { border-color: #6c5ce7; }
.toolbar-right button.primary { background: #6c5ce7; border-color: #6c5ce7; color: #fff; }
.toolbar-right button.primary:disabled { opacity: 0.4; cursor: not-allowed; }
.toolbar-right input[type=file] { display: none; }

.main { flex: 1; display: flex; overflow: hidden; }
.canvas-panel {
  flex: 1; overflow: auto; display: flex;
  justify-content: center; align-items: flex-start; padding: 12px;
  position: relative;
}
.canvas-wrap { position: relative; display: inline-block; }
.canvas-wrap canvas { display: block; }
.canvas-wrap canvas.base { position: relative; }
.canvas-wrap canvas.overlay {
  position: absolute; top: 0; left: 0;
  cursor: crosshair; opacity: 0.4;
}

.sidebar {
  width: 290px; background: #15151e; border-left: 1px solid #2a2a3a;
  display: flex; flex-direction: column; flex-shrink: 0; overflow-y: auto;
}
.section { padding: 10px 12px; border-bottom: 1px solid #2a2a3a; }
.section h3 { font-size: 11px; color: #8888a0; margin-bottom: 6px; text-transform: uppercase; letter-spacing: 0.5px; }

.region-list { padding: 6px; max-height: 180px; overflow-y: auto; }
.region-item {
  display: flex; align-items: center; gap: 5px;
  padding: 3px 6px; font-size: 11px; font-family: monospace;
  border-radius: 3px; cursor: pointer;
}
.region-item:hover { background: #1c1c2a; }
.region-item .num { color: #6c5ce7; width: 18px; }
.region-item .del { color: #fd79a8; cursor: pointer; margin-left: auto; padding: 0 4px; }
.region-item .del:hover { color: #ff5252; }

.wm-field { margin-bottom: 6px; }
.wm-field label { display: block; font-size: 11px; color: #8888a0; margin-bottom: 2px; }
.wm-field input, .wm-field select {
  width: 100%; padding: 4px 6px; background: #1c1c2a;
  border: 1px solid #2a2a3a; border-radius: 3px; color: #e0e0f0; font-size: 12px;
}
.wm-field input:focus, .wm-field select:focus { outline: none; border-color: #6c5ce7; }
.wm-row { display: flex; gap: 6px; }
.wm-row .wm-field { flex: 1; }
input[type=range] { width: 100%; height: 4px; -webkit-appearance: none; background: #2a2a3a; border-radius: 2px; }
input[type=range]::-webkit-slider-thumb { -webkit-appearance: none; width: 14px; height: 14px; border-radius: 50%; background: #6c5ce7; cursor: pointer; }
.range-label { display: flex; justify-content: space-between; font-size: 10px; color: #8888a0; }

.status-bar {
  padding: 5px 12px; background: #15151e; border-top: 1px solid #2a2a3a;
  font-size: 11px; color: #8888a0;
}
.toast {
  position: fixed; bottom: 50px; left: 50%; transform: translateX(-50%);
  background: #6c5ce7; color: #fff; padding: 7px 18px;
  border-radius: 5px; font-size: 12px; z-index: 200;
  display: none; box-shadow: 0 4px 16px rgba(108,92,231,0.4);
}
.toast.error { background: #fd79a8; }
.empty-msg { position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);color:#444;font-size:20px;text-align:center;pointer-events:none; }
</style>
</head>
<body>

<div class="header">
  <div class="tab-bar">
    <div class="tab active" onclick="switchTab('blur')" id="tabBlur">🔒 Blur</div>
    <div class="tab" onclick="switchTab('word')" id="tabWord">📄 Word</div>
    <div class="tab" onclick="switchTab('form')" id="tabForm">📝 Form</div>
  </div>
  <div class="toolbar">
    <div class="toolbar-left">
      <label for="fileInput" class="btn">📁 Open PDF</label>
      <input type="file" id="fileInput" accept=".pdf">
    </div>
    <div class="toolbar-right" id="toolbarBlur">
      <button onclick="analyzePDF()" id="analyzeBtn" disabled>🔍 Analyze</button>
      <button onclick="processPDF()" class="primary" id="processBtn" disabled>⚡ Process</button>
      <button onclick="downloadPDF()" id="downloadBtn" disabled>⬇ Download</button>
    </div>
    <div class="toolbar-right" id="toolbarWord" style="display:none">
      <button onclick="previewPage()" id="previewBtn">👁 Preview</button>
      <button onclick="convertToWord()" class="primary" id="convertBtn" disabled>📄 Download Word</button>
    </div>
    <div class="toolbar-right" id="toolbarForm" style="display:none">
      <button onclick="previewForm()" id="formPreviewBtn">👁 Preview</button>
      <button onclick="downloadForm()" class="primary" id="formDlBtn" disabled>⬇ Download Fillable PDF</button>
    </div>
  </div>
</div>

<div class="main">
  <div class="canvas-panel" id="canvasPanel">
    <div class="canvas-wrap" id="canvasWrap">
      <canvas id="baseCanvas" class="base"></canvas>
      <canvas id="overlayCanvas" class="overlay"></canvas>
    </div>
    <div class="empty-msg" id="emptyMsg">📄<br>Upload a PDF</div>
  </div>

  <div class="sidebar" id="sidebar" style="display:none">
    <!-- Pages (shown in both modes) -->
    <div class="section" id="pagesSection">
      <h3>📄 Pages</h3>
      <div style="display:flex;gap:3px;flex-wrap:wrap;margin-top:4px" id="pageThumbs"></div>
    </div>

    <!-- Blur controls (shown in blur mode) -->
    <div id="sidebarBlur">
      <div class="section">
        <h3>🔍 Blur</h3>
        <p style="font-size:11px;color:#666;margin-bottom:4px">Drag on the image to add regions</p>
        <div style="margin-bottom:4px">
          <div class="range-label"><span>Intensity</span><span id="blurVal">25</span></div>
          <input type="range" id="blurIntensity" min="5" max="50" value="25">
        </div>
        <div class="region-list" id="regionList">
          <div style="color:#555;font-size:11px;text-align:center;padding:10px">No regions yet</div>
        </div>
      </div>

      <!-- Watermark -->
      <div class="section">
        <h3>💧 Watermark <span onclick="clearWatermark()" style="color:#fd79a8;cursor:pointer;font-size:12px;float:right">✖ Clear</span></h3>
      <div class="wm-field">
        <label>Text</label>
        <input type="text" id="wmText" value="ESGov">
      </div>
      <div class="wm-field">
        <label>Size: <span id="wmSizeVal" style="color:#6c5ce7">80</span></label>
        <input type="range" id="wmSize" min="20" max="200" value="80">
      </div>
      <div class="wm-field">
        <label>Spacing: <span id="wmSpacingVal" style="color:#6c5ce7">2.0</span></label>
        <input type="range" id="wmSpacing" min="0.5" max="5.0" step="0.1" value="2.0">
      </div>
      <div class="wm-row">
        <div class="wm-field">
          <label>Opacity</label>
          <select id="wmOpacity">
            <option value="10">10%</option>
            <option value="20">20%</option>
            <option value="30" selected>30%</option>
            <option value="50">50%</option>
            <option value="70">70%</option>
          </select>
        </div>
        <div class="wm-field">
          <label>Color</label>
          <div style="display:flex;gap:2px;flex-wrap:wrap">
            <span onclick="document.getElementById('wmColor').value='#888'" style="display:inline-block;width:18px;height:18px;border-radius:3px;background:#888;cursor:pointer;border:2px solid transparent"></span>
            <span onclick="document.getElementById('wmColor').value='#f00'" style="display:inline-block;width:18px;height:18px;border-radius:3px;background:#f00;cursor:pointer;border:2px solid transparent"></span>
            <span onclick="document.getElementById('wmColor').value='#00f'" style="display:inline-block;width:18px;height:18px;border-radius:3px;background:#00f;cursor:pointer;border:2px solid transparent"></span>
            <span onclick="document.getElementById('wmColor').value='#0a0'" style="display:inline-block;width:18px;height:18px;border-radius:3px;background:#0a0;cursor:pointer;border:2px solid transparent"></span>
            <input type="color" id="wmColor" value="#888" style="width:22px;height:18px;padding:0;border:1px solid #2a2a3a;background:#1c1c2a;cursor:pointer;border-radius:3px">
          </div>
        </div>
      </div>
    </div>

    <!-- PDF → Word controls -->
    <div id="sidebarWord" style="display:none">
      <div class="section">
        <h3>📄 Word</h3>
        <p style="font-size:12px;color:#aaa;margin-top:4px;line-height:1.5">
          Convert PDF to editable Word document.
        </p>
        <button onclick="convertToWord()" class="btn primary" style="width:100%;margin-top:8px" id="convertBtn2" disabled>📄 Download Word</button>
      </div>
    </div>

    <!-- Fillable Form controls -->
    <div id="sidebarForm" style="display:none">
      <div class="section">
        <h3>📝 Form Fields</h3>
        <p style="font-size:11px;color:#888;margin-bottom:6px">Click on the document image to place a fillable text field.</p>
        <button onclick="addFormField()" class="btn" style="width:100%;margin-bottom:6px">➕ Add Field</button>
        <div class="region-list" id="formFieldList" style="max-height:300px">
          <div style="color:#555;font-size:11px;text-align:center;padding:10px">No fields yet. Click "Add Field" then click the document.</div>
        </div>
      </div>
      <div class="section">
        <h3>⬇ Download</h3>
        <button onclick="downloadForm()" class="btn primary" style="width:100%" id="formDlBtn2" disabled>📄 Download Fillable PDF</button>
      </div>
    </div>
  </div>
</div>

<div class="status-bar"><span id="statusText">Open a PDF to get started</span></div>
<div class="toast" id="toast"></div>

<script>
const state = { sessionId: null, pages: [], currentPage: 0, regions: {}, scale: 1 };
const baseCanvas = document.getElementById('baseCanvas');
const overlayCanvas = document.getElementById('overlayCanvas');
const overlayCtx = overlayCanvas.getContext('2d');
let isDrawing = false, dx = 0, dy = 0;

document.getElementById('fileInput').addEventListener('change', async e => {
  const f = e.target.files[0]; if (!f) return;
  const fd = new FormData(); fd.append('pdf', f);
  try {
    const r = await fetch('/upload', { method:'POST', body:fd });
    const d = await r.json();
    if (!r.ok) throw new Error(d.error);
    state.sessionId = d.session_id; state.pages = d.pages;
    state.currentPage = 0; state.regions = {}; detectedBlocks = {};
    document.getElementById('emptyMsg').style.display = 'none';
    document.getElementById('sidebar').style.display = 'flex';
    document.getElementById('processBtn').disabled = false;
    document.getElementById('analyzeBtn').disabled = false;
    document.getElementById('convertBtn').disabled = false;
    document.getElementById('convertBtn2').disabled = false;
    document.getElementById('formDlBtn').disabled = false;
    document.getElementById('formDlBtn2').disabled = false;
    renderThumbs(); renderPage(0);
    status(d.pages.length + ' pages loaded');
    toast('PDF loaded: ' + f.name);
  } catch(e) { toast('Error: ' + e.message, true); }
});

function renderPage(idx) {
  const page = state.pages[idx];
  if (!page) return;
  state.currentPage = idx;
  const panel = document.getElementById('canvasPanel');
  const s = Math.min((panel.clientWidth-24)/page.w, (panel.clientHeight-24)/page.h, 1.5);
  state.scale = s;
  baseCanvas.width = page.w; baseCanvas.height = page.h;
  baseCanvas.style.width = (page.w*s)+'px'; baseCanvas.style.height = (page.h*s)+'px';
  overlayCanvas.width = page.w; overlayCanvas.height = page.h;
  overlayCanvas.style.width = (page.w*s)+'px'; overlayCanvas.style.height = (page.h*s)+'px';
  const img = new Image();
  img.onload = () => { baseCanvas.getContext('2d').drawImage(img,0,0); redrawOverlay(); if (Object.keys(detectedBlocks).length) renderDetectedBlocks(); };
  img.src = page.dataUrl;
  status('Page '+(idx+1)+' of '+state.pages.length);
  updateRegionList();
  updateThumbs();
}

function renderThumbs() {
  const c = document.getElementById('pageThumbs');
  c.innerHTML = '';
  state.pages.forEach((p,i) => {
    const t = document.createElement('div');
    t.style.cssText = 'width:36px;height:48px;border:2px solid #2a2a3a;border-radius:3px;overflow:hidden;cursor:pointer;background-size:cover;background-position:center;background-image:url('+p.dataUrl+')';
    t.onclick = () => renderPage(i);
    t.dataset.idx = i;
    c.appendChild(t);
  });
}
function updateThumbs() {
  document.getElementById('pageThumbs').childNodes.forEach(t => {
    t.style.borderColor = parseInt(t.dataset.idx) === state.currentPage ? '#6c5ce7' : '#2a2a3a';
  });
}

// Region drawing
overlayCanvas.addEventListener('mousedown', e => {
  const r = overlayCanvas.getBoundingClientRect(), s = state.scale;
  dx = (e.clientX - r.left) / s; dy = (e.clientY - r.top) / s;
  isDrawing = true;
});
overlayCanvas.addEventListener('mousemove', e => {
  if (!isDrawing) return;
  const r = overlayCanvas.getBoundingClientRect(), s = state.scale;
  const cx = (e.clientX - r.left) / s, cy = (e.clientY - r.top) / s;
  const x = Math.min(dx,cx), y = Math.min(dy,cy), w = Math.abs(cx-dx), h = Math.abs(cy-dy);
  redrawOverlay();
  overlayCtx.strokeStyle = '#6c5ce7'; overlayCtx.lineWidth = 2/s;
  overlayCtx.setLineDash([4/s,4/s]);
  overlayCtx.strokeRect(x,y,w,h);
  overlayCtx.fillStyle = 'rgba(108,92,231,0.15)'; overlayCtx.fillRect(x,y,w,h);
});
overlayCanvas.addEventListener('mouseup', e => {
  if (!isDrawing) return; isDrawing = false; overlayCtx.setLineDash([]);
  const r = overlayCanvas.getBoundingClientRect(), s = state.scale;
  const ex = (e.clientX-r.left)/s, ey = (e.clientY-r.top)/s;
  const x = Math.min(dx,ex), y = Math.min(dy,ey), w = Math.abs(ex-dx), h = Math.abs(ey-dy);
  if (w < 10 || h < 10) { redrawOverlay(); return; }
  const pi = state.currentPage;
  if (!state.regions[pi]) state.regions[pi] = [];
  state.regions[pi].push({id:Date.now(), x:Math.round(x), y:Math.round(y), w:Math.round(w), h:Math.round(h)});
  redrawOverlay(); updateRegionList();
});

function redrawOverlay() {
  const s = state.scale;
  overlayCtx.clearRect(0, 0, overlayCanvas.width, overlayCanvas.height);
  const rs = state.regions[state.currentPage] || [];
  rs.forEach((r,i) => {
    overlayCtx.strokeStyle = '#ffd740'; overlayCtx.lineWidth = 2/s;
    overlayCtx.strokeRect(r.x, r.y, r.w, r.h);
    overlayCtx.fillStyle = 'rgba(255,215,64,0.12)'; overlayCtx.fillRect(r.x, r.y, r.w, r.h);
    overlayCtx.fillStyle = '#ffd740';
    overlayCtx.font = Math.max(9, 12/s)+'px monospace';
    overlayCtx.fillText('#'+(i+1), r.x+3/s, r.y+11/s);
  });
}

function updateRegionList() {
  const rs = state.regions[state.currentPage] || [];
  const el = document.getElementById('regionList');
  if (!rs.length) {
    el.innerHTML = '<div style="color:#555;font-size:11px;text-align:center;padding:8px">Drag on image to add</div>';
    return;
  }
  el.innerHTML = rs.map((r,i) =>
    '<div class="region-item"><span class="num">#'+(i+1)+'</span><span>('+r.x+','+r.y+') '+r.w+'\u00d7'+r.h+'</span><span class="del" onclick="delRegion('+i+')">\u2716</span></div>'
  ).join('');
}

function delRegion(i) {
  const pi = state.currentPage;
  if (state.regions[pi]) state.regions[pi].splice(i,1);
  redrawOverlay(); updateRegionList();
}

// Process (blur + watermark)
async function processPDF() {
  const btn = document.getElementById('processBtn');
  btn.disabled = true; btn.textContent = '⏳ Processing...';
  try {
    const intensity = parseInt(document.getElementById('blurIntensity').value);
    const wmList = [];
    const text = document.getElementById('wmText').value.trim();
    if (text) {
      wmList.push({
        text: text,
        size: parseInt(document.getElementById('wmSize').value),
        spacing: parseFloat(document.getElementById('wmSpacing').value),
        opacity: parseInt(document.getElementById('wmOpacity').value),
        color: document.getElementById('wmColor').value,
      });
    }
    const r = await fetch('/process', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({
        session_id: state.sessionId,
        regions: state.regions,
        radius: intensity,
        watermarks: wmList
      })
    });
    const d = await r.json();
    if (!r.ok) throw new Error(d.error);
    state.pages = d.pages;
    renderPage(state.currentPage);
    document.getElementById('downloadBtn').disabled = false;
    toast('✅ Processed!');
  } catch(e) { toast('Error: '+e.message, true); }
  finally { btn.disabled = false; btn.textContent = '⚡ Process'; }
}

// Analyze PDF — detect text blocks
let detectedBlocks = {};

async function analyzePDF() {
  const btn = document.getElementById('analyzeBtn');
  btn.disabled = true; btn.textContent = '⏳ Analyzing...';
  try {
    const r = await fetch('/analyze', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({session_id: state.sessionId})
    });
    const d = await r.json();
    if (!r.ok) throw new Error(d.error);
    detectedBlocks = d.blocks || {};
    renderDetectedBlocks();
    const total = Object.values(detectedBlocks).reduce((a,b) => a + b.length, 0);
    toast('🔍 ' + total + ' text blocks detected. Click a block to add to blur.');
    status(total + ' text blocks found');
  } catch(e) { toast('Error: '+e.message, true); }
  finally { btn.disabled = false; btn.textContent = '🔍 Analyze'; }
}

function renderDetectedBlocks() {
  const s = state.scale;
  const ctx = overlayCanvas.getContext('2d');
  const blocks = detectedBlocks[state.currentPage] || [];
  // Redraw blur regions first
  redrawOverlay();
  // Then draw detected blocks
  ctx.globalAlpha = 0.3;
  blocks.forEach((b, i) => {
    ctx.fillStyle = '#00ff88';
    ctx.fillRect(b.x, b.y, b.w, b.h);
    ctx.strokeStyle = '#00ff88';
    ctx.lineWidth = 1.5/s;
    ctx.setLineDash([]);
    ctx.strokeRect(b.x, b.y, b.w, b.h);
    ctx.globalAlpha = 0.8;
    ctx.fillStyle = '#00ff88';
    ctx.font = Math.max(8, 10/s)+'px monospace';
    ctx.fillText('📄'+(i+1), b.x+2/s, b.y+10/s);
    ctx.globalAlpha = 0.3;
  });
  ctx.globalAlpha = 1.0;
  // Make blocks clickable via overlay click
}

// Intercept overlay click for detected blocks
overlayCanvas.addEventListener('click', e => {
  if (isDrawing) return;
  const r = overlayCanvas.getBoundingClientRect(), s = state.scale;
  const cx = (e.clientX - r.left) / s, cy = (e.clientY - r.top) / s;
  const blocks = detectedBlocks[state.currentPage] || [];
  for (let i = blocks.length - 1; i >= 0; i--) {
    const b = blocks[i];
    if (cx >= b.x && cx <= b.x + b.w && cy >= b.y && cy <= b.y + b.h) {
      // Add to blur regions
      const pi = state.currentPage;
      if (!state.regions[pi]) state.regions[pi] = [];
      state.regions[pi].push({id: Date.now(), x: Math.round(b.x), y: Math.round(b.y), w: Math.round(b.w), h: Math.round(b.h)});
      // Remove from detected blocks so it doesn't get added again
      blocks.splice(i, 1);
      redrawOverlay();
      renderDetectedBlocks();
      updateRegionList();
      toast('➕ Block added to blur');
      break;
    }
  }
});

async function downloadPDF() {
  try {
    const r = await fetch('/download', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({session_id: state.sessionId})
    });
    if (!r.ok) throw new Error('Download failed');
    const blob = await r.blob();
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = 'processed_output.pdf';
    a.click();
    URL.revokeObjectURL(a.href);
    toast('✅ Downloaded!');
  } catch(e) { toast('Error: '+e.message, true); }
}

document.getElementById('blurIntensity').addEventListener('input', function() {
  document.getElementById('blurVal').textContent = this.value;
});

document.getElementById('wmSize').addEventListener('input', function() {
  document.getElementById('wmSizeVal').textContent = this.value;
});
document.getElementById('wmSpacing').addEventListener('input', function() {
  document.getElementById('wmSpacingVal').textContent = parseFloat(this.value).toFixed(1);
});

function clearWatermark() {
  document.getElementById('wmText').value = '';
  document.getElementById('wmSize').value = 80;
  document.getElementById('wmSizeVal').textContent = '80';
  document.getElementById('wmSpacing').value = 2.0;
  document.getElementById('wmSpacingVal').textContent = '2.0';
  document.getElementById('wmOpacity').value = '30';
  document.getElementById('wmColor').value = '#888';
}

function toast(msg, err) {
  const el = document.getElementById('toast');
  el.textContent = msg; el.className = 'toast'+(err?' error':'');
  el.style.display = 'block';
  clearTimeout(el._timer);
  el._timer = setTimeout(() => el.style.display = 'none', 2500);
}
function status(msg) { document.getElementById('statusText').textContent = msg; }

// Tab switching
let activeTab = 'blur';
function switchTab(tab) {
  activeTab = tab;
  document.getElementById('tabBlur').className = 'tab' + (tab === 'blur' ? ' active' : '');
  document.getElementById('tabWord').className = 'tab' + (tab === 'word' ? ' active' : '');
  document.getElementById('tabForm').className = 'tab' + (tab === 'form' ? ' active' : '');
  document.getElementById('toolbarBlur').style.display = tab === 'blur' ? '' : 'none';
  document.getElementById('toolbarWord').style.display = tab === 'word' ? '' : 'none';
  document.getElementById('toolbarForm').style.display = tab === 'form' ? '' : 'none';
  document.getElementById('sidebarBlur').style.display = tab === 'blur' ? '' : 'none';
  document.getElementById('sidebarWord').style.display = tab === 'word' ? '' : 'none';
  document.getElementById('sidebarForm').style.display = tab === 'form' ? '' : 'none';
  // Show/hide overlay interactivity
  overlayCanvas.style.cursor = (tab === 'blur' || tab === 'form') ? 'crosshair' : 'default';
  overlayCanvas.style.opacity = (tab === 'blur' || tab === 'form') ? '0.4' : '0';
  if (tab === 'form' && state.pages.length) redrawFormFields();
}

// PDF → Word conversion
async function convertToWord() {
  const btn = document.getElementById('convertBtn');
  const btn2 = document.getElementById('convertBtn2');
  btn.disabled = true; btn.textContent = '⏳...';
  if (btn2) { btn2.disabled = true; btn2.textContent = '⏳...'; }
  try {
    const r = await fetch('/to-word', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({session_id: state.sessionId})
    });
    if (!r.ok) { const d = await r.json(); throw new Error(d.error || 'Conversion failed'); }
    const blob = await r.blob();
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = 'converted.docx';
    a.click();
    URL.revokeObjectURL(a.href);
    toast('✅ Word downloaded!');
  } catch(e) { toast('Error: '+e.message, true); }
  finally {
    btn.disabled = false; btn.textContent = '📄 Download Word';
    if (btn2) { btn2.disabled = false; btn2.textContent = '📄 Download Word'; }
  }
}

// 👁 Preview — show the current page in the canvas (reuse renderPage)
function previewPage() {
  if (!state.pages.length) return;
  renderPage(state.currentPage);
  toast('📄 Preview: Page ' + (state.currentPage+1));
}

// ═══════════════════════════════════════
// 📝 Fillable Form
// ═══════════════════════════════════════
let formFields = {};  // { pageIndex: [{id, x, y, w, h, label, value}, ...] }
let fieldIdCounter = 0;
let addingField = false;

function addFormField() {
  if (!state.pages.length) return;
  addingField = !addingField;
  overlayCanvas.style.cursor = addingField ? 'copy' : 'crosshair';
  toast(addingField ? 'Click on the document to place a field' : 'Add field cancelled');
}

// Override overlay mousedown to handle form field placement
const origMouseDown = overlayCanvas.onmousedown;
overlayCanvas.addEventListener('mousedown', function(e) {
  if (activeTab !== 'form') return;
  if (addingField) {
    const r = overlayCanvas.getBoundingClientRect(), s = state.scale;
    const cx = (e.clientX - r.left) / s, cy = (e.clientY - r.top) / s;
    const pi = state.currentPage;
    if (!formFields[pi]) formFields[pi] = [];
    const fw = 120, fh = 20;
    formFields[pi].push({
      id: ++fieldIdCounter,
      x: Math.round(cx - fw/2), y: Math.round(cy - fh/2),
      w: fw, h: fh,
      label: 'Field ' + fieldIdCounter,
      value: ''
    });
    redrawFormFields();
    updateFormFieldList();
    addingField = false;
    overlayCanvas.style.cursor = 'crosshair';
    toast('✅ Field added. Drag corners to resize.');
    return;
  }
  // Check if clicking on existing field corner to resize
});

function redrawFormFields() {
  const ctx = overlayCanvas.getContext('2d');
  const s = state.scale;
  ctx.clearRect(0, 0, overlayCanvas.width, overlayCanvas.height);
  const fs = formFields[state.currentPage] || [];
  fs.forEach((f, i) => {
    // Draw field box
    ctx.fillStyle = 'rgba(108,92,231,0.08)';
    ctx.fillRect(f.x, f.y, f.w, f.h);
    ctx.strokeStyle = '#6c5ce7';
    ctx.lineWidth = 1.5/s;
    ctx.setLineDash([4/s, 2/s]);
    ctx.strokeRect(f.x, f.y, f.w, f.h);
    ctx.setLineDash([]);
    // Draw label
    ctx.fillStyle = '#6c5ce7';
    ctx.font = Math.max(8, 10/s)+'px sans-serif';
    ctx.fillText(f.label, f.x+2/s, f.y-2/s);
    // Draw value if exists
    if (f.value) {
      ctx.fillStyle = '#e0e0f0';
      ctx.font = Math.max(9, 11/s)+'px sans-serif';
      ctx.fillText(f.value, f.x+3/s, f.y+f.h/2+4/s);
    }
  });
}

function updateFormFieldList() {
  const fs = formFields[state.currentPage] || [];
  const el = document.getElementById('formFieldList');
  if (!fs.length) {
    el.innerHTML = '<div style="color:#555;font-size:11px;text-align:center;padding:10px">No fields yet. Click "Add Field" then click the document.</div>';
    return;
  }
  el.innerHTML = fs.map((f, i) =>
    '<div class="region-item" onclick="editFieldLabel('+f.id+')">' +
    '<span class="num">#'+(i+1)+'</span>' +
    '<span style="flex:1;overflow:hidden;text-overflow:ellipsis">'+(f.label)+'</span>' +
    '<input id="fld_'+f.id+'" value="'+f.value.replace(/"/g,'&quot;')+'" ' +
    'style="width:60px;padding:1px 3px;background:#1c1c2a;border:1px solid #2a2a3a;border-radius:2px;color:#e0e0f0;font-size:10px" ' +
    'onchange="formFieldValueChanged('+f.id+',this.value)" placeholder="value">' +
    '<span class="del" onclick="delFormField('+f.id+')">\u2716</span></div>'
  ).join('');
}

function editFieldLabel(fid) {
  const f = findField(fid);
  if (!f) return;
  const newLabel = prompt('Field label:', f.label);
  if (newLabel && newLabel.trim()) {
    f.label = newLabel.trim();
    redrawFormFields();
    updateFormFieldList();
  }
}

function formFieldValueChanged(fid, val) {
  const f = findField(fid);
  if (f) { f.value = val; redrawFormFields(); }
}

function delFormField(fid) {
  const pi = state.currentPage;
  if (formFields[pi]) {
    formFields[pi] = formFields[pi].filter(f => f.id !== fid);
    redrawFormFields();
    updateFormFieldList();
  }
}

function findField(fid) {
  for (const pi in formFields) {
    for (const f of formFields[pi]) {
      if (f.id === fid) return f;
    }
  }
  return null;
}

// Preview form — show filled values on the overlay
function previewForm() {
  if (!state.pages.length) return;
  renderPage(state.currentPage);
  redrawFormFields();
  toast('👁 Preview: showing form fields');
}

// Download fillable PDF
async function downloadForm() {
  const btn = document.getElementById('formDlBtn');
  const btn2 = document.getElementById('formDlBtn2');
  btn.disabled = true; btn.textContent = '⏳...';
  if (btn2) { btn2.disabled = true; btn2.textContent = '⏳...'; }
  try {
    const r = await fetch('/download-form', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({
        session_id: state.sessionId,
        fields: formFields
      })
    });
    if (!r.ok) { const d = await r.json(); throw new Error(d.error || 'Download failed'); }
    const blob = await r.blob();
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = 'fillable_form.pdf';
    a.click();
    URL.revokeObjectURL(a.href);
    toast('✅ Fillable PDF downloaded!');
  } catch(e) { toast('Error: '+e.message, true); }
  finally {
    btn.disabled = false; btn.textContent = '⬇ Download Fillable PDF';
    if (btn2) { btn2.disabled = false; btn2.textContent = '📄 Download Fillable PDF'; }
  }
}

window.addEventListener('resize', () => { if (state.pages.length) renderPage(state.currentPage); });
</script>
</body>
</html>
'''


# ══════════════════════════════════════════════════════════════
# PDF → Word conversion
# ══════════════════════════════════════════════════════════════

def pdf_to_word(pdf_path, output_path):
    """Convert PDF to an editable Word document using pdf2docx engine.
    Falls back to image embedding for scanned/image-based PDFs."""
    import shutil

    # Pre-check: if the PDF uses OCR/hidden fonts, it's a scanned document
    # Skip directly to image fallback for better visual quality
    pdf_check = fitz.open(pdf_path)
    is_scanned = False
    ocr_fonts = {'hiddenhorzocr', 'hiddenvertocr', 'ocra', 'ocrb'}
    for page_num in range(min(len(pdf_check), 3)):  # Check first 3 pages
        page = pdf_check.load_page(page_num)
        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        font = span.get("font", "").lower().replace(" ", "")
                        if any(ocr in font for ocr in ocr_fonts):
                            is_scanned = True
                            break
                if is_scanned:
                    break
        if is_scanned:
            break
    pdf_check.close()

    if is_scanned:
        # Scanned PDF: render pages as images and embed in Word
        doc = Document()
        pdf = fitz.open(pdf_path)
        for page_num in range(len(pdf)):
            page = pdf.load_page(page_num)
            rect = page.rect
            if page_num > 0:
                doc.add_section()
            section = doc.sections[-1]
            section.page_width = Pt(rect.width)
            section.page_height = Pt(rect.height)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img_path = os.path.join(os.path.dirname(output_path), f'_page{page_num}.png')
            with open(img_path, 'wb') as f:
                f.write(img_bytes)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Pt(rect.width))
            os.remove(img_path)
        pdf.close()
        for p in doc.paragraphs[:]:
            if not p.text.strip() and not p.runs:
                try:
                    p._element.getparent().remove(p._element)
                except:
                    pass
        doc.save(output_path)
        return

    # Text-based PDF: use pdf2docx for proper conversion
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
    except Exception as e:
        # Fallback to image embedding
        doc = Document()
        pdf = fitz.open(pdf_path)
        for page_num in range(len(pdf)):
            page = pdf.load_page(page_num)
            rect = page.rect
            if page_num > 0:
                doc.add_section()
            section = doc.sections[-1]
            section.page_width = Pt(rect.width)
            section.page_height = Pt(rect.height)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img_path = os.path.join(os.path.dirname(output_path), f'_page{page_num}.png')
            with open(img_path, 'wb') as f:
                f.write(img_bytes)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Pt(rect.width))
            os.remove(img_path)
        pdf.close()
        for p in doc.paragraphs[:]:
            if not p.text.strip() and not p.runs:
                try:
                    p._element.getparent().remove(p._element)
                except:
                    pass
        doc.save(output_path)


# ══════════════════════════════════════════════════════════════
# Routes
# ══════════════════════════════════════════════════════════════

@app.route('/')
def index():
    return HTML


@app.route('/upload', methods=['POST'])
def upload():
    if 'pdf' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['pdf']
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Only PDF files accepted'}), 400
    session_id = uuid.uuid4().hex[:12]
    paths = get_session_paths(session_id)
    file.save(paths['pdf'])
    images = pdf_to_images(paths['pdf'])
    pages = []
    for i, img in enumerate(images):
        buf = io.BytesIO()
        img.save(buf, 'PNG')
        pages.append({
            'index': i,
            'dataUrl': f'data:image/png;base64,{base64.b64encode(buf.getvalue()).decode()}',
            'w': img.width, 'h': img.height
        })
    # Save meta
    with open(paths['meta'], 'w') as f:
        _json.dump({'original_name': file.filename}, f)

    # Save originals
    originals_dir = paths['originals']
    if os.path.exists(originals_dir):
        import shutil
        shutil.rmtree(originals_dir)
    os.makedirs(originals_dir, exist_ok=True)
    for i, img in enumerate(images):
        img.save(os.path.join(originals_dir, f'{i}.png'))

    # Save initial preview copies to processed dir
    processed_dir = paths['processed_dir']
    if os.path.exists(processed_dir):
        import shutil
        shutil.rmtree(processed_dir)
    os.makedirs(processed_dir, exist_ok=True)
    for i, img in enumerate(images):
        img.save(os.path.join(processed_dir, f'{i}.png'))
    return jsonify({'session_id': session_id, 'pages': pages})


@app.route('/analyze', methods=['POST'])
def analyze():
    """Detect text blocks in the PDF."""
    data = request.json
    sid = data.get('session_id')
    paths = get_session_paths(sid)
    if not os.path.exists(paths['pdf']):
        return jsonify({'error': 'Session not found'}), 404

    doc = fitz.open(paths['pdf'])
    blocks_per_page = {}
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        # Get text blocks with coordinates
        text_blocks = page.get_text("blocks")
        blocks = []
        for b in text_blocks:
            if b[6] == 0:  # type 0 = text block
                x0, y0, x1, y1 = b[0], b[1], b[2], b[3]
                # Scale by 2x to match image coordinates
                x0 *= 2; y0 *= 2; x1 *= 2; y1 *= 2
                w = x1 - x0
                h = y1 - y0
                if w > 10 and h > 5:  # skip tiny fragments
                    blocks.append({
                        'x': round(x0), 'y': round(y0),
                        'w': round(w), 'h': round(h),
                        'text': b[4][:80]  # first 80 chars for reference
                    })
        blocks_per_page[str(page_num)] = blocks
    doc.close()
    return jsonify({'blocks': blocks_per_page})


@app.route('/process', methods=['POST'])
def process():
    data = request.json
    sid = data.get('session_id')
    paths = get_session_paths(sid)
    if not os.path.exists(paths['pdf']):
        return jsonify({'error': 'Session not found'}), 404
    radius = data.get('radius', 25)
    regions = data.get('regions', {})
    watermarks = data.get('watermarks', [])

    # Load original images (always from originals, so re-process doesn't stack)
    originals_dir = paths['originals']
    if not os.path.exists(originals_dir):
        return jsonify({'error': 'Session not found'}), 404

    images = []
    for f in sorted(os.listdir(originals_dir), key=lambda x: int(x.split('.')[0])):
        if f.endswith('.png'):
            images.append(Image.open(os.path.join(originals_dir, f)))

    if not images:
        return jsonify({'error': 'No images found'}), 404

    processed_dir = paths['processed_dir']
    os.makedirs(processed_dir, exist_ok=True)
    processed_images = []
    for pi, img in enumerate(images):
        img = img.copy()
        pr = regions.get(str(pi), [])
        if pr:
            img = apply_blur(img, pr, radius)
        for wm in watermarks:
            if wm.get('text'):
                img = apply_watermark(img, wm)
        img.save(os.path.join(processed_dir, f'{pi}.png'))
        processed_images.append(img)

    # Return previews
    pages = []
    for i, img in enumerate(processed_images):
        buf = io.BytesIO()
        img.save(buf, 'PNG')
        pages.append({
            'index': i,
            'dataUrl': f'data:image/png;base64,{base64.b64encode(buf.getvalue()).decode()}',
            'w': img.width, 'h': img.height
        })
    return jsonify({'pages': pages})


@app.route('/download', methods=['POST'])
def download():
    data = request.json
    sid = data.get('session_id')
    paths = get_session_paths(sid)
    if not os.path.exists(paths['pdf']):
        return jsonify({'error': 'Session not found'}), 404

    # Load images from processed dir (always the latest version)
    processed_dir = paths['processed_dir']
    images = []
    if os.path.exists(processed_dir):
        for f in sorted(os.listdir(processed_dir), key=lambda x: int(x.split('.')[0])):
            if f.endswith('.png'):
                images.append(Image.open(os.path.join(processed_dir, f)))

    if not images:
        return jsonify({'error': 'No images found'}), 404

    # Get original filename for download name
    orig_name = 'processed_output.pdf'
    try:
        with open(paths['meta']) as f:
            meta = _json.load(f)
            orig_name = meta.get('original_name', 'processed_output.pdf')
            if orig_name.lower().endswith('.pdf'):
                orig_name = orig_name[:-4] + '_blur_watermark.pdf'
            else:
                orig_name = orig_name + '_blur_watermark.pdf'
    except:
        pass

    output = os.path.join(paths['processed_dir'], 'output.pdf')
    images_to_pdf(images, output)
    return send_file(output, as_attachment=True,
                     download_name=orig_name,
                     mimetype='application/pdf')


@app.route('/to-word', methods=['POST'])
def to_word():
    """Convert uploaded PDF to editable Word document."""
    data = request.json
    sid = data.get('session_id')
    paths = get_session_paths(sid)
    if not os.path.exists(paths['pdf']):
        return jsonify({'error': 'Session not found'}), 404

    # Get original filename for download name
    orig_name = 'converted.docx'
    try:
        with open(paths['meta']) as f:
            meta = _json.load(f)
            orig_name = meta.get('original_name', 'document.docx')
            if orig_name.lower().endswith('.pdf'):
                orig_name = orig_name[:-4] + '.docx'
            else:
                orig_name = orig_name + '.docx'
    except:
        pass

    try:
        output_path = os.path.join(paths['processed_dir'], 'converted.docx')
        pdf_to_word(paths['pdf'], output_path)
        if not os.path.exists(output_path):
            return jsonify({'error': 'Conversion failed - no output file'}), 500
        return send_file(output_path, as_attachment=True,
                         download_name=orig_name,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'error': f'Conversion error: {str(e)}'}), 500


@app.route('/download-form', methods=['POST'])
def download_form():
    """Generate a fillable PDF form with AcroForm fields overlaid on the scanned document."""
    data = request.json
    sid = data.get('session_id')
    fields_data = data.get('fields', {})
    paths = get_session_paths(sid)
    if not os.path.exists(paths['pdf']):
        return jsonify({'error': 'Session not found'}), 404

    try:
        # Open the original PDF to get page dimensions
        src_pdf = fitz.open(paths['pdf'])
        num_pages = len(src_pdf)

        # Create output PDF
        out_pdf = fitz.open()

        for page_num in range(num_pages):
            src_page = src_pdf.load_page(page_num)
            rect = src_page.rect

            # Render page as image at 2x
            mat = fitz.Matrix(2.0, 2.0)
            pix = src_page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            # Create new page with same dimensions
            new_page = out_pdf.new_page(width=rect.width, height=rect.height)

            # Insert the rendered page image as background
            img_rect = fitz.Rect(0, 0, rect.width, rect.height)
            new_page.insert_image(img_rect, stream=img_bytes)

            # Add AcroForm text fields for this page
            page_fields = fields_data.get(str(page_num), [])
            for f in page_fields:
                widget = fitz.Widget()
                widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT  # 7
                widget.field_name = f.get('label', f'field_{f.get("id", 0)}')
                widget.rect = fitz.Rect(
                    f['x'], f['y'],
                    f['x'] + f['w'],
                    f['y'] + f['h']
                )
                widget.text_font = "Helv"
                widget.text_font_size = 11
                widget.border_color = (0.42, 0.36, 0.91)  # #6c5ce7
                widget.fill_color = (1, 1, 1)
                widget.text_color = (0, 0, 0)
                widget.border_width = 1
                if f.get('value'):
                    widget.field_value = f['value']
                    widget.text_value = f['value']
                new_page.add_widget(widget)

        src_pdf.close()

        # Save to temp path
        output_path = os.path.join(paths['processed_dir'], 'fillable_form.pdf')
        out_pdf.save(output_path, garbage=4, deflate=True)
        out_pdf.close()

        if not os.path.exists(output_path):
            return jsonify({'error': 'Form generation failed'}), 500

        return send_file(output_path, as_attachment=True,
                         download_name='fillable_form.pdf',
                         mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': f'Form error: {str(e)}'}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8777))
    print('╔══════════════════════════════════════╗')
    print('║  PDF Blur + Watermark Tool           ║')
    print(f'║  http://0.0.0.0:{port:<39}║')
    print('╚══════════════════════════════════════╝')
    app.run(host='0.0.0.0', port=port, debug=True)
